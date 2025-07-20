from docx import Document

# Crear el documento Word formato ATS con enlaces funcionales
doc = Document()

# Título principal
doc.add_heading('JHON ALEXANDER VALENCIA MARULANDA', 0)

# Información de contacto con enlaces funcionales
doc.add_paragraph(
    "Dirección: Calle 64 b No 20ª55 Torre 2, apto 301, Mirador de la Palma, Manizales\n"
    "Teléfono Fijo: (601) 8896780 | Móvil: (57) 311 3398444\n"
    "Correo electrónico: jhonalexanderv@gmail.com\n"
    "Fecha de nacimiento: 20/07/1975 | C.C: 75.078.042 | T.P: CL230-31983\n"
    "Portafolio: https://studio--engineerview.us-central1.hosted.app\n"
    "GitHub: https://github.com/Jhonalex75\n"
    "Certificación (Coursera): https://www.coursera.org/account/accomplishments/specialization/certificate/4YEAW3TB8D78\n"
    "Publicación científica: http://www.sciencedirect.com/science/article/pii/S0304386X07002101"
)

# Perfil Profesional
doc.add_heading("Perfil Profesional", level=1)
doc.add_paragraph(
    "Ingeniero Mecánico con maestría en Ciencias de la Ingeniería Mención Procesamiento de Minerales. Con experiencia en construcciones metal mecánicas, incluyendo el pre-diseño, presupuesto, fabricación, montaje, mantenimiento y administración general. "
    "Conocimiento en logística y desarrollo de soluciones tecnológicas, evaluación e innovación de elementos y sistemas mecánicos. Trabajo en equipo, instrucción, toma de decisiones, aprendizaje continuo, creatividad, y manejo de relaciones interpersonales."
)

# Formación Académica
doc.add_heading("Formación Académica", level=1)
doc.add_paragraph(
    "• Magíster en Ciencias de la Ingeniería, Mención Procesamiento de Minerales – Universidad de Antofagasta, Chile (2003-2006)\n"
    "• Introduction to Project Management Principles – University of California, Irvine (2015)\n"
    "• Ingeniero Mecánico y de Manufacturas – Universidad Autónoma de Manizales, Colombia (1994-1999)\n"
    "• Estudios primarios y secundarios – Colegio Mayor de Nuestra Señora, Manizales (1982-1992)"
)

# Otros Estudios
doc.add_heading("Otros Estudios", level=1)
doc.add_paragraph(
    "• SENA: Técnicas de Automatización y Control, Programación Básica Fresadora CNC, Mantenimiento de Sistemas Neumáticos, Programa Supervisor\n"
    "• Universidad de Antofagasta: Innovación y Desarrollo de Productos Químicos y Bioproductos"
)

# Experiencia Laboral
doc.add_heading("Experiencia Laboral", level=1)
experiences = [
    "Mina El Gran Porvenir – Jefe de mantenimiento, selección de equipos, montaje, producción (Oct 2024 – Abr 2025)",
    "EGM Colombia S.A.S – Jefe de mantenimiento mecánico (May 2024 – Sep 2024)",
    "Consorcio PTAP Tibitoc 20 – Especialista Mecánico (Dic 2023 – Abr 2024)",
    "Schrader Camargo – Coordinador de Ingeniería y Presupuesto (May 2023 – Dic 2023)",
    "O-I Peldar – Ingeniero Mecánico Senior (Abr 2022 – May 2023)",
    "China Harbour Engineering Company – Ingeniero Mecánico (Feb 2022 – Abr 2022)",
    "CIM – Líder de Obra (Mar 2021 – Ene 2022)",
    "Merit Consultants Intl – Coordinador Montaje Mecánico y Estructural (Nov 2018 – Nov 2020)",
    "Gran Colombia Gold – Jefe Departamento Técnico (Dic 2016 – Nov 2018)",
    "Consorcio HHA Aguas de Aburrá – Ingeniero Montaje Mecánico (Sep 2014 – Dic 2016)",
    "TEPSA S.A – Ingeniero de Proyectos (Mar 2011 – Sep 2014)",
    "Drill Metal Ingeniería – Mantenimiento e instalaciones para múltiples empresas (2008 – 2010)",
    "Pioneros Ingeniería – Montajes ACASA (Mar – Jul 2008)",
    "Mastil Ingeniería – Montaje gas natural y refrigeración (2007)",
    "Universidad de Antofagasta – Docente de Instrumentación Industrial (2002 – 2006)",
    "Hidromec Ltda – Supervisor área Hidráulica (2005 – 2006)",
    "Minera Escondida – Inspección de tuberías (2004)",
    "Asiterm – Aislamiento térmico en turbinas (2004)",
    "Conymet Dunlop-Simsa – Jefe Control y Calidad (2002 – 2003)"
]
for exp in experiences:
    doc.add_paragraph(f"• {exp}")

# Idiomas
doc.add_heading("Idiomas", level=1)
doc.add_paragraph("Inglés 70% hablado y escrito – Centro Colombo Americano")

# Trabajos Académicos
doc.add_heading("Trabajos Académicos", level=1)
doc.add_paragraph(
    "• Tesis de Magister: Saltpeter extraction and modelling of caliche mineral heap leaching\n"
    "• Tesis Pregrado: Diseño de Software CAD/CAM para Engranajes Rectos (Tesis Meritoria)"
)

# Habilidades y Conocimientos
doc.add_heading("Habilidades y Conocimientos", level=1)
doc.add_paragraph(
    "• Microsoft Office: Word, Excel, PowerPoint, Outlook, MS Project\n"
    "• Programación: Matlab, VBA, Python\n"
    "• Software CAD: SolidWorks, Inventor, AutoCAD Plant 3D\n"
    "• Trabajo en equipo, liderazgo, diseño y supervisión de proyectos"
)

# Referencias Personales
doc.add_heading("Referencias Personales", level=1)
doc.add_paragraph(
    "• Dr. Luís Alberto Cisternas – lcisternas@uantof.cl | Tel: (55)-637323\n"
    "• Alexander Valencia – Ing. Mecánico | CEL: 3206888811\n"
    "• Carlos Roure – Ing. Mecánico | CEL: 3002161814\n"
    "• Dereck Seah – Ing. Mecánico | +59 171325296"
)

# Referencias Laborales
doc.add_heading("Referencias Laborales", level=1)
doc.add_paragraph(
    "• Franklin Gonzalez – Procurement Manager, Continental Goldcorp | ingfgg@gmail.com | +34 642414997\n"
    "• Nicholas Blanchette – General Manager, Merit Consultants Intl | Nick.blanchette@cementation.com | +57 3114174118"
)

# Guardar el archivo
doc_path = "Hoja_de_Vida_Jhon_Valencia_ATS_Enlaces.docx"
doc.save(doc_path)

doc_path
